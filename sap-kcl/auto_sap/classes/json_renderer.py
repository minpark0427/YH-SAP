"""
JSON-based SAP renderer — builds docx from structured JSON output.

Instead of docxtpl (which inserts plain text into tagged positions),
this renderer uses python-docx to create properly formatted paragraphs
from JSON-structured LLM output.

Expected JSON format per section:
{
    "paragraphs": [
        {"text": "...", "style": "body"},
        {"text": "...", "style": "body"},
        {"text": "item 1", "style": "bullet"},
        {"text": "item 2", "style": "bullet"},
    ]
}

Styles:
  "body" → Normal paragraph
  "bullet" → List Paragraph / bullet
  "bold" → Bold paragraph (for sub-labels within a section)
  "table" → Table (with "headers" and "rows" fields instead of "text")
"""

import json
import re
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_llm_json(text: str) -> list[dict]:
    """Parse LLM output as JSON. Falls back to plain text if JSON parsing fails."""
    text = text.strip()

    # Try to extract JSON from the response
    # LLM might wrap it in ```json ... ``` or have extra text
    json_match = re.search(r'\{[\s\S]*"paragraphs"[\s\S]*\}', text)
    if json_match:
        try:
            data = json.loads(json_match.group())
            if "paragraphs" in data:
                return data["paragraphs"]
        except json.JSONDecodeError:
            pass

    # Try to parse as JSON array directly
    array_match = re.search(r'\[[\s\S]*\]', text)
    if array_match:
        try:
            data = json.loads(array_match.group())
            if isinstance(data, list) and len(data) > 0:
                return data
        except json.JSONDecodeError:
            pass

    # Fallback: convert plain text to paragraph list
    return _plain_text_to_paragraphs(text)


def _is_skip_marker(text: str) -> tuple[bool, str]:
    """Check if text contains a [SKIP: reason] marker."""
    import re
    match = re.search(r'\[SKIP:\s*(.+?)\]', text)
    if match:
        return True, match.group(1).strip()
    # Also catch common variations
    if text.strip().startswith("[SKIP") or "not specified in the protocol" in text.lower()[:100]:
        return True, text.strip()
    return False, ""


def _plain_text_to_paragraphs(text: str) -> list[dict]:
    """Convert plain text (possibly with markdown) to paragraph dicts."""
    # Strip markdown
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'(?<!\w)\*([^*\n]+?)\*(?!\w)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)

    paragraphs = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if re.match(r'^[-•]\s+', line):
            paragraphs.append({"text": re.sub(r'^[-•]\s+', '', line), "style": "bullet"})
        elif re.match(r'^\d+\.\s+', line):
            paragraphs.append({"text": line, "style": "body"})
        else:
            paragraphs.append({"text": line, "style": "body"})

    return paragraphs


class JsonSapRenderer:
    """Render SAP sections from JSON-structured LLM output into a docx."""

    def __init__(self, template_path: str | Path):
        """Initialize with the Yuhan template docx as base."""
        self.template_path = Path(template_path)

    def render(self, sap_content: dict[str, str], output_path: str | Path,
               protocol_tables: dict[str, list[dict]] | None = None) -> None:
        """Render all SAP sections into the template docx.

        Args:
            sap_content: dict of tag → LLM response (JSON string or plain text)
            output_path: where to save the rendered docx
            protocol_tables: optional dict from ProtocolTableExtractor.extract_all()
                           Tables are injected into relevant SAP sections.
        """
        doc = Document(str(self.template_path))

        # Find and replace each {{ tag }} paragraph with formatted content
        paragraphs_to_process = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if '{{' in text and '}}' in text:
                # Extract tag name
                match = re.search(r'\{\{\s*(\w+)\s*\}\}', text)
                if match:
                    tag = match.group(1)
                    if tag in sap_content:
                        paragraphs_to_process.append((i, p, tag))

        # Inject protocol tables into relevant SAP sections
        if protocol_tables:
            sap_content = self._inject_protocol_tables(sap_content, protocol_tables)

        # Process in reverse to preserve indices
        for i, para, tag in reversed(paragraphs_to_process):
            content = sap_content[tag]
            if content == "ERROR" or not content.strip():
                # Leave error marker
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = "[ERROR: Content generation failed]"
                continue

            parsed = parse_llm_json(content)
            self._replace_paragraph_with_content(doc, para, parsed)

        # Handle the statistical_software_version special case
        for p in doc.paragraphs:
            if '{{' in p.text and 'statistical_software_version' in p.text:
                version = sap_content.get('statistical_software_version', '9.4 (or later)')
                # Just do simple text replacement
                for run in p.runs:
                    if '{{' in run.text:
                        run.text = run.text.replace(
                            '{{ statistical_software_version }}', version
                        ).replace(
                            '{{statistical_software_version}}', version
                        )

        # Post-process
        self._fix_blue_text(doc)
        self._remove_instruction_boxes(doc)

        # Save
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        print(f"SAP saved to {output_path}")

    def _replace_paragraph_with_content(self, doc, placeholder_para, content_items: list[dict]):
        """Replace a placeholder paragraph with formatted content paragraphs and tables."""
        parent = placeholder_para._element.getparent()
        placeholder_elem = placeholder_para._element

        last_elem = placeholder_elem
        for item in content_items:
            style = item.get("style", "body")

            # Handle table items
            if style == "table":
                headers = item.get("headers", [])
                rows = item.get("rows", [])
                caption = item.get("caption", "")
                if headers or rows:
                    tbl_elem = self._create_table_element(headers, rows, caption)
                    last_elem.addnext(tbl_elem)
                    last_elem = tbl_elem
                continue

            text = item.get("text", "").strip()
            if not text:
                continue

            # Check for SKIP marker
            is_skip, skip_reason = _is_skip_marker(text)

            new_p = copy.deepcopy(placeholder_elem)
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            for child in list(new_p):
                if child.tag == qn('w:r'):
                    new_p.remove(child)

            pPr = new_p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                new_p.insert(0, pPr)

            if style == "bullet":
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    pPr.append(pStyle)
                pStyle.set(qn('w:val'), 'ListParagraph')
            else:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    pPr.remove(pStyle)

            new_r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            if is_skip:
                # Yellow highlight for skipped sections
                highlight = OxmlElement('w:highlight')
                highlight.set(qn('w:val'), 'yellow')
                rPr.append(highlight)
                # Italic for skip text
                rPr.append(OxmlElement('w:i'))
                text = f"[REVIEW NEEDED] {skip_reason}"
            elif style == "bold":
                rPr.append(OxmlElement('w:b'))

            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)

            color = OxmlElement('w:color')
            color.set(qn('w:val'), '000000')
            rPr.append(color)

            new_r.append(rPr)

            new_t = OxmlElement('w:t')
            new_t.text = text
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            new_p.append(new_r)

            last_elem.addnext(new_p)
            last_elem = new_p

        parent.remove(placeholder_elem)

    def _create_table_element(self, headers: list[str], rows: list[list[str]], caption: str = "") -> OxmlElement:
        """Create a docx table XML element from headers and rows."""
        all_rows = [headers] + rows if headers else rows
        if not all_rows:
            return OxmlElement('w:p')

        n_cols = max(len(r) for r in all_rows)

        tbl = OxmlElement('w:tbl')

        # Table properties
        tblPr = OxmlElement('w:tblPr')
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        tblPr.append(tblW)
        # Add borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        # Table grid
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(n_cols):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), '1440')  # ~1 inch per column
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        # Rows
        for r_idx, row_data in enumerate(all_rows):
            tr = OxmlElement('w:tr')
            for c_idx in range(n_cols):
                tc = OxmlElement('w:tc')
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')

                # Header row: bold + smaller font
                sz = OxmlElement('w:sz')
                if r_idx == 0 and headers:
                    rPr.append(OxmlElement('w:b'))
                    sz.set(qn('w:val'), '16')  # 8pt header
                else:
                    sz.set(qn('w:val'), '16')  # 8pt body
                rPr.append(sz)
                r.append(rPr)

                t = OxmlElement('w:t')
                t.text = row_data[c_idx] if c_idx < len(row_data) else ""
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        return tbl

    # Table-to-SAP tag mapping: which protocol tables go into which SAP sections
    TABLE_TO_SAP_MAP = {
        "objectives": ["objectives"],           # Objectives & Endpoints tables
        "overall_design": ["introduction"],      # Overall design → introduction
        "schedule_of_activities": ["visit_windows"],  # SoA → visit windows
        "investigational_products": ["ip_exposure"],   # IP info → exposure section
        "pk_parameters": ["primary_efficacy"],   # PK params → primary analysis
        "lab_assessments": ["lab_parameters"],   # Lab tests → lab parameters
    }

    def _inject_protocol_tables(self, sap_content: dict[str, str],
                                 protocol_tables: dict[str, list[dict]]) -> dict[str, str]:
        """Inject protocol tables into SAP content as table items in JSON."""
        sap_content = dict(sap_content)  # don't mutate original

        for table_key, sap_tags in self.TABLE_TO_SAP_MAP.items():
            tables = protocol_tables.get(table_key, [])
            if not tables:
                continue

            for sap_tag in sap_tags:
                if sap_tag not in sap_content:
                    continue

                # Parse existing content
                existing = sap_content[sap_tag]
                parsed = parse_llm_json(existing)

                # Append protocol tables
                for t in tables:
                    # Add caption paragraph before table
                    parsed.append({
                        "text": f"[Protocol Table: {t['caption']}]",
                        "style": "bold",
                    })
                    parsed.append({
                        "style": "table",
                        "headers": t["headers"],
                        "rows": t["rows"],
                        "caption": t["caption"],
                    })

                # Re-serialize
                sap_content[sap_tag] = json.dumps({"paragraphs": parsed})

        return sap_content

    def _fix_blue_text(self, doc):
        """Convert all blue text to black throughout the document."""
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.color and run.font.color.rgb:
                    if str(run.font.color.rgb) == '0000FF':
                        run.font.color.rgb = RGBColor(0, 0, 0)

    def _remove_instruction_boxes(self, doc):
        """Remove gray instruction boxes (tables with 'Instruction' in first cell)."""
        tables_to_remove = []
        for table in doc.tables:
            try:
                first_cell = table.rows[0].cells[0].text.strip()
                if first_cell.startswith("Instruction") or first_cell.startswith("Overall Instruction"):
                    tables_to_remove.append(table)
            except (IndexError, AttributeError):
                continue

        for table in tables_to_remove:
            table._element.getparent().remove(table._element)

        if tables_to_remove:
            print(f"  Removed {len(tables_to_remove)} instruction boxes from template")
