"""Run SAP generation 5 times and validate each output."""
import subprocess, os, sys, time, re

os.chdir("/Users/mymini/Vibecoding/YH-SAP/sap-kcl")
sys.path.insert(0, ".")

PROTOCOL = "../YH00000-101_Clinical Protocol_Ver1.0_22Feb2024.docx"
PYTHON = "/Users/mymini/Vibecoding/YH-SAP/YH-SAP-venv/bin/python3"

results = []

for run_num in range(1, 6):
    output_dir = f"SAPs_run{run_num}"
    os.makedirs(output_dir, exist_ok=True)

    print(f"\n{'='*60}")
    print(f"RUN {run_num}/5")
    print(f"{'='*60}", flush=True)

    t0 = time.time()
    result = subprocess.run(
        [PYTHON, "-m", "WriteSAPs.yuhan_sap_json", PROTOCOL, "--test", "--output-dir", output_dir],
        capture_output=True, text=True, timeout=300
    )
    elapsed = round(time.time() - t0)

    if result.returncode != 0:
        print(f"  CRASH ({elapsed}s): {result.stderr[:200]}")
        results.append(("CRASH", run_num, elapsed))
        continue

    # Find files
    content_file = docx_file = None
    for f in os.listdir(output_dir):
        if f.endswith("_content.txt"): content_file = os.path.join(output_dir, f)
        elif f.endswith(".docx"): docx_file = os.path.join(output_dir, f)

    if not content_file or not docx_file:
        print(f"  NO OUTPUT ({elapsed}s)")
        results.append(("NO_OUTPUT", run_num, elapsed))
        continue

    # Validate
    from docx import Document
    from docx.oxml.ns import qn
    with open(content_file) as f: content = f.read()
    doc = Document(docx_file)

    from auto_sap.prompts.prompts_yuhan_v1 import PROMPTS_DICTIONARY
    missing = [t for t in PROMPTS_DICTIONARY if f"{t}:" not in content]
    errors = content.count("ERROR")
    skips = len(re.findall(r'\[SKIP:', content))
    instr = sum(1 for t in doc.tables if t.rows[0].cells[0].text.strip().startswith("Instruction"))
    md = sum(1 for p in doc.paragraphs if p.text.strip().startswith("# ") or "**" in p.text)
    blue = sum(1 for p in doc.paragraphs for r in p.runs
               if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == "0000FF")
    yellow = 0
    for p in doc.paragraphs:
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                hl = rPr.find(qn('w:highlight'))
                if hl is not None and hl.get(qn('w:val')) == 'yellow': yellow += 1
    size = os.path.getsize(docx_file)

    checks = {
        "tags_46": len(missing) == 0,
        "errors_0": errors == 0,
        "no_instruction": instr == 0,
        "no_markdown": md == 0,
        "no_blue": blue == 0,
        "has_tables": len(doc.tables) > 5,
        "docx_valid": size > 100000,
        "has_yellow": yellow > 0,
        "has_skips": skips > 0,
    }
    failed = [k for k, v in checks.items() if not v]
    status = "ALL_PASS" if not failed else "FAIL"

    print(f"  {status} | {elapsed}s | paras={len(doc.paragraphs)} | tables={len(doc.tables)} | {size//1024}KB | skips={skips} | yellows={yellow}")
    if failed: print(f"  FAILED: {failed}")
    results.append((status, run_num, elapsed))
    sys.stdout.flush()

print(f"\n{'='*60}")
print("FINAL SUMMARY")
print(f"{'='*60}")
for status, num, elapsed in results:
    print(f"  Run {num}: {status} ({elapsed}s)")
pass_count = sum(1 for s, _, _ in results if s == "ALL_PASS")
print(f"\nPassed: {pass_count}/5")
if pass_count == 5:
    print("VERDICT: PIPELINE STABLE — all 5 runs passed")
else:
    print(f"VERDICT: {5-pass_count} failures — needs investigation")
