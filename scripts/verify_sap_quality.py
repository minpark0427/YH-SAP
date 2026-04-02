"""
SAP Quality Verification — Section-by-section comparison of AI SAP vs Human SAP.

Automated gap detection + LLM-based content evaluation.
Produces a structured report with scores and actionable improvement items.

Usage:
    python scripts/verify_sap_quality.py \
        --ai-sap sap-kcl/SAPs_verify/YH_SAP_YH00000-101_Clinical_Protocol_Ver1.0_22Feb2024_content.txt \
        --human-sap "YH00000-101_SAP ver1.0(20240827).docx" \
        --protocol "YH00000-101_Clinical Protocol_Ver1.0_22Feb2024.docx" \
        --output reports/verification_report.md
"""

import argparse
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


# ============================================================================
# Section Mapping: Human SAP (other company) ↔ AI SAP (Yuhan template)
# ============================================================================

SECTION_MAP = [
    # (ai_tag, human_section_id, human_section_title, category)
    ("abbreviations", "Abbreviations", "LIST OF ABBREVIATIONS", "structure"),
    ("introduction", "1", "INTRODUCTION", "content"),
    ("objectives", "2", "TRIAL OBJECTIVES", "content"),
    ("general_considerations", "4.1", "General Considerations for Data Summarization and Analysis", "content"),
    ("screened_set", "4.2.Screened", "Screened Set", "content"),
    ("randomized_set", "4.2.Randomized", "Randomized Set", "content"),
    ("safety_set", "4.2.Safety", "Safety Analysis Set", "content"),
    ("fas_definition", "4.2.FAS", "Full Analysis Set", "content"),
    ("pps_definition", "4.2.PPS", "Per-Protocol Set (not in human SAP)", "content"),
    ("subject_disposition", "4.3", "Subject Disposition", "content"),
    ("demographics", "4.5", "Demographics and Other Baseline Characteristics", "content"),
    ("ip_exposure", "4.12", "IP Compliance (for Part B)", "content"),
    ("treatment_compliance", "4.12", "IP Compliance (for Part B)", "content"),
    ("medical_history", "4.10", "Medical History", "content"),
    ("concomitant_medication", "4.11", "Prior and Concomitant Medication", "content"),
    ("special_tests", "4.6-4.9", "Breathalyzer/Hepatitis/HIV/Syphilis/Urine Drug", "content"),
    ("primary_efficacy", "4.15", "Pharmacokinetics Analyses", "critical"),
    ("secondary_efficacy", "4.16", "Pharmacodynamic Analyses", "critical"),
    ("additional_efficacy", "4.17", "Exploratory Analyses", "content"),
    ("preliminary_pk", "4.19", "Preliminary PK Analyses", "content"),
    ("preliminary_pd", "4.20", "Preliminary PD Analyses", "content"),
    ("blinded_analyses", "4.21", "Blinded Analyses", "content"),
    ("adverse_events", "4.14.AE", "Adverse Events", "content"),
    ("lab_parameters", "4.14.Lab", "Clinical Laboratory Parameters", "content"),
    ("vital_signs", "4.14.VS", "Vital Signs", "content"),
    ("ecg", "4.14.ECG", "12-lead Electrocardiogram", "content"),
    ("physical_exam", "4.14.PE", "Physical Examination", "content"),
    ("other_safety", "4.14.Other", "Other Safety", "content"),
    ("interim_analysis", "N/A", "Interim Analysis (not in human SAP)", "content"),
    ("dmc_analysis", "N/A", "DMC Analysis (not in human SAP)", "content"),
    ("sample_size", "3.6", "Sample Size Determination", "critical"),
    ("statistical_software_version", "5", "SOFTWARE LIST", "content"),
    ("statistical_considerations", "4.1+", "Statistical Considerations (spread across 4.1)", "content"),
    ("visit_windows", "N/A", "Visit Time Windows (not explicitly in human SAP)", "content"),
    ("derived_variables", "N/A", "Derived Variables (not explicitly in human SAP)", "content"),
    ("repeated_assessments", "4.1", "General Considerations (repeated assessments)", "content"),
    ("missing_last_dose_date", "N/A", "Missing Last Dose Date (not in human SAP)", "content"),
    ("missing_ae_severity", "N/A", "Missing AE Severity (not in human SAP)", "content"),
    ("missing_causality", "N/A", "Missing Causality (not in human SAP)", "content"),
    ("missing_dates", "N/A", "Missing Dates (not in human SAP)", "content"),
    ("lab_char_values", "N/A", "Lab Character Values (not in human SAP)", "content"),
    ("protocol_deviation", "4.4", "Protocol Deviation", "content"),
    ("baseline_definition", "4.13", "Definition of Baseline", "critical"),
    ("subgroup_analysis", "4.18", "Subgroup Analysis", "content"),
    ("references", "6", "REFERENCE LIST", "structure"),
    ("protocol_changes", "N/A", "Protocol Changes", "content"),
]


@dataclass
class SectionScore:
    tag: str
    human_section: str
    category: str  # structure, content, critical
    # Automated metrics
    ai_word_count: int = 0
    has_skip: bool = False
    skip_reasons: list = field(default_factory=list)
    human_has_content: bool = False
    human_word_count: int = 0
    # LLM evaluation scores (0-5)
    protocol_fidelity: int = -1  # Does AI content match protocol?
    completeness: int = -1       # Are all protocol details captured?
    specificity: int = -1        # Formulas, versions, exact numbers present?
    comparison_to_human: int = -1  # How does AI compare to human SAP?
    # Qualitative
    gaps: list = field(default_factory=list)
    strengths: list = field(default_factory=list)
    recommendations: list = field(default_factory=list)


# ============================================================================
# Extraction Helpers
# ============================================================================

def extract_ai_sections(content_path: str) -> dict[str, str]:
    """Parse the AI SAP content file into {tag: text} dict."""
    with open(content_path) as f:
        content = f.read()

    sections = {}
    # Find all tag markers
    tag_positions = []
    for m in re.finditer(r'^([a-z_]+): ', content, re.MULTILINE):
        tag_positions.append((m.group(1), m.start()))

    for i, (tag, start) in enumerate(tag_positions):
        end = tag_positions[i + 1][1] if i + 1 < len(tag_positions) else len(content)
        sections[tag] = content[start + len(tag) + 2:end].strip()

    return sections


def extract_human_sections(docx_path: str) -> dict[str, str]:
    """Extract body content from human SAP docx (skip TOC, tables)."""
    doc = Document(docx_path)

    full_text = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name
        if style.startswith('toc') or style.startswith('TOC'):
            continue
        if 'Heading' in style:
            full_text.append(f"\n=== {text} ===\n")
        else:
            full_text.append(text)

    return "\n".join(full_text)


def count_tables(docx_path: str) -> int:
    doc = Document(docx_path)
    return len(doc.tables)


def extract_table_headers(docx_path: str) -> list[str]:
    """Extract first row of each table for comparison."""
    doc = Document(docx_path)
    headers = []
    for i, t in enumerate(doc.tables):
        if t.rows:
            cells = [c.text.strip()[:40] for c in t.rows[0].cells]
            headers.append(f"Table {i+1}: {' | '.join(cells[:4])}")
    return headers


# ============================================================================
# Automated Checks (no LLM needed)
# ============================================================================

def run_automated_checks(ai_sections: dict[str, str], human_text: str,
                         ai_docx_path: str, human_docx_path: str) -> dict:
    """Run all automated quality checks."""

    results = {
        "skip_markers": {},
        "word_counts": {},
        "table_comparison": {},
        "coverage": {},
    }

    # 1. SKIP marker analysis
    total_skips = 0
    for tag, text in ai_sections.items():
        skips = re.findall(r'\[SKIP:([^\]]+)\]', text)
        if skips:
            results["skip_markers"][tag] = [s.strip() for s in skips]
            total_skips += len(skips)
    results["skip_markers"]["_total"] = total_skips

    # 2. Word count per section
    for tag, text in ai_sections.items():
        results["word_counts"][tag] = len(text.split())

    # 3. Table comparison
    ai_tables = count_tables(ai_docx_path)
    human_tables = count_tables(human_docx_path)
    results["table_comparison"] = {
        "ai_tables": ai_tables,
        "human_tables": human_tables,
        "gap": human_tables - ai_tables,
        "ai_table_headers": extract_table_headers(ai_docx_path),
        "human_table_headers": extract_table_headers(human_docx_path),
    }

    # 4. Coverage: sections with <30 words (likely incomplete)
    thin_sections = {t: wc for t, wc in results["word_counts"].items() if wc < 30}
    results["coverage"]["thin_sections"] = thin_sections
    results["coverage"]["total_sections"] = len(ai_sections)
    results["coverage"]["sections_with_skip"] = len(results["skip_markers"]) - 1  # exclude _total
    results["coverage"]["total_words"] = sum(results["word_counts"].values())

    return results


# ============================================================================
# LLM-based Evaluation
# ============================================================================

def claude_evaluate(prompt: str, model: str = "claude-sonnet-4-5") -> str:
    """Call claude CLI for evaluation."""
    env = os.environ.copy()
    env.pop("ANTHROPIC_API_KEY", None)
    try:
        result = subprocess.run(
            ["claude", "--print", "-p", "-", "--model", model],
            capture_output=True, text=True, input=prompt, env=env, timeout=300,
        )
        if result.returncode != 0:
            return f"ERROR: {result.stderr[:300]}"
        return result.stdout.strip()
    except subprocess.TimeoutExpired:
        return "ERROR: Timeout"
    except FileNotFoundError:
        return "ERROR: claude CLI not found"


def evaluate_section_pair(tag: str, ai_text: str, human_text_excerpt: str,
                          category: str, human_section_title: str) -> dict:
    """LLM-evaluate one AI SAP section against human SAP."""

    prompt = f"""You are a senior biostatistician reviewing a Statistical Analysis Plan (SAP) section.

TASK: Compare the AI-generated section against the human-written reference, focusing on whether
the AI SAP accurately captures protocol content. The two SAPs use DIFFERENT templates, so
section structure will differ — compare the CONTENT and MEANING, not the format.

AI SAP SECTION: {tag}
HUMAN SAP SECTION: {human_section_title}
IMPORTANCE: {category.upper()}

=== AI-GENERATED CONTENT ===
{ai_text[:6000]}

=== HUMAN-WRITTEN REFERENCE (relevant excerpt) ===
{human_text_excerpt[:6000]}

Score each dimension 0-5 (0=missing, 1=poor, 2=below average, 3=adequate, 4=good, 5=excellent):

Respond in EXACTLY this JSON format (no markdown, no code blocks):
{{
  "protocol_fidelity": <0-5>,
  "completeness": <0-5>,
  "specificity": <0-5>,
  "comparison_to_human": <0-5>,
  "gaps": ["gap1", "gap2"],
  "strengths": ["strength1"],
  "recommendations": ["rec1"],
  "summary": "one sentence overall assessment"
}}

SCORING GUIDELINES:
- protocol_fidelity: Does AI content accurately reflect what the protocol says? (5=perfect, 0=fabricated)
- completeness: Are all relevant protocol details captured? ([SKIP] on items the human SAP covers = low score)
- specificity: Are formulas, version numbers, exact timepoints, exact values present? (vague = low)
- comparison_to_human: Overall, does AI section match or exceed human quality? (5=better, 3=equivalent, 0=far worse)

Return ONLY the JSON object. No other text."""

    response = claude_evaluate(prompt)

    # Parse JSON from response
    try:
        # Try to extract JSON even if wrapped in markdown
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            return json.loads(json_match.group())
    except (json.JSONDecodeError, AttributeError):
        pass

    return {
        "protocol_fidelity": -1,
        "completeness": -1,
        "specificity": -1,
        "comparison_to_human": -1,
        "gaps": [f"Failed to parse LLM response: {response[:200]}"],
        "strengths": [],
        "recommendations": [],
        "summary": "Evaluation failed",
    }


def find_human_excerpt(human_text: str, section_title: str, tag: str) -> str:
    """Extract the relevant portion of human SAP text for comparison."""
    # Search for the section heading in human text
    patterns = [
        section_title,
        section_title.split("(")[0].strip(),  # Without parenthetical
        tag.replace("_", " "),
    ]

    best_start = -1
    for pattern in patterns:
        idx = human_text.lower().find(pattern.lower())
        if idx >= 0:
            best_start = max(0, idx - 100)
            break

    if best_start < 0:
        return "(No matching section found in human SAP)"

    # Extract ~3000 chars around the match
    return human_text[best_start:best_start + 3000]


# ============================================================================
# Report Generation
# ============================================================================

def generate_report(auto_results: dict, section_scores: list[dict],
                    ai_content_path: str, human_docx_path: str) -> str:
    """Generate the full markdown verification report."""

    lines = [
        "# SAP Quality Verification Report",
        "",
        f"**AI SAP**: `{ai_content_path}`",
        f"**Human SAP**: `{human_docx_path}`",
        f"**Generated**: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "---",
        "",
        "## 1. Executive Summary",
        "",
    ]

    # Compute aggregate scores
    scored = [s for s in section_scores if s.get("protocol_fidelity", -1) >= 0]
    if scored:
        avg_fidelity = sum(s["protocol_fidelity"] for s in scored) / len(scored)
        avg_completeness = sum(s["completeness"] for s in scored) / len(scored)
        avg_specificity = sum(s["specificity"] for s in scored) / len(scored)
        avg_comparison = sum(s["comparison_to_human"] for s in scored) / len(scored)
        overall = (avg_fidelity + avg_completeness + avg_specificity + avg_comparison) / 4

        lines.extend([
            f"| Metric | Score (0-5) |",
            f"|--------|-------------|",
            f"| Protocol Fidelity | {avg_fidelity:.1f} |",
            f"| Completeness | {avg_completeness:.1f} |",
            f"| Specificity | {avg_specificity:.1f} |",
            f"| Comparison to Human | {avg_comparison:.1f} |",
            f"| **Overall** | **{overall:.1f}** |",
            "",
        ])

    # Skip markers summary
    skip_total = auto_results["skip_markers"].get("_total", 0)
    skip_sections = auto_results["coverage"]["sections_with_skip"]
    total_sections = auto_results["coverage"]["total_sections"]
    lines.extend([
        f"- **[SKIP] markers**: {skip_total} across {skip_sections}/{total_sections} sections",
        f"- **Total AI word count**: {auto_results['coverage']['total_words']:,}",
        f"- **Tables**: AI={auto_results['table_comparison']['ai_tables']}, "
        f"Human={auto_results['table_comparison']['human_tables']} "
        f"(gap: {auto_results['table_comparison']['gap']})",
        "",
    ])

    # 2. SKIP marker analysis
    lines.extend([
        "---",
        "",
        "## 2. [SKIP] Marker Analysis",
        "",
        "These are sections where the AI decided protocol content was insufficient.",
        "Items marked ⚠️ indicate the human SAP DID provide content — meaning the AI should too.",
        "",
    ])

    # Known sections where human SAP has content but AI skipped
    human_has_content = {
        "fas_definition", "baseline_definition", "protocol_deviation",
        "statistical_software_version", "medical_history", "concomitant_medication",
        "visit_windows", "derived_variables", "missing_dates", "missing_ae_severity",
        "missing_causality", "missing_last_dose_date", "lab_char_values",
        "preliminary_pk", "preliminary_pd", "blinded_analyses",
    }

    for tag, reasons in sorted(auto_results["skip_markers"].items()):
        if tag == "_total":
            continue
        flag = " ⚠️ **Human SAP has content**" if tag in human_has_content else ""
        lines.append(f"### `{tag}`{flag}")
        for r in reasons:
            lines.append(f"- {r}")
        lines.append("")

    # 3. Section-by-section scores
    lines.extend([
        "---",
        "",
        "## 3. Section-by-Section Evaluation",
        "",
        "| Tag | Words | Fidelity | Complete | Specific | vs Human | Summary |",
        "|-----|-------|----------|----------|----------|----------|---------|",
    ])

    for s in section_scores:
        tag = s["tag"]
        wc = auto_results["word_counts"].get(tag, 0)
        f = s.get("protocol_fidelity", -1)
        c = s.get("completeness", -1)
        sp = s.get("specificity", -1)
        h = s.get("comparison_to_human", -1)
        summary = s.get("summary", "")[:60]

        def fmt(v):
            if v < 0:
                return "—"
            if v <= 1:
                return f"🔴 {v}"
            if v <= 2:
                return f"🟡 {v}"
            if v <= 3:
                return f"🟢 {v}"
            return f"✅ {v}"

        lines.append(f"| `{tag}` | {wc} | {fmt(f)} | {fmt(c)} | {fmt(sp)} | {fmt(h)} | {summary} |")

    lines.append("")

    # 4. Detailed findings per section
    lines.extend([
        "---",
        "",
        "## 4. Detailed Findings",
        "",
    ])

    for s in section_scores:
        lines.append(f"### `{s['tag']}`")
        if s.get("gaps"):
            lines.append("**Gaps:**")
            for g in s["gaps"]:
                lines.append(f"- {g}")
        if s.get("strengths"):
            lines.append("**Strengths:**")
            for st in s["strengths"]:
                lines.append(f"- {st}")
        if s.get("recommendations"):
            lines.append("**Recommendations:**")
            for r in s["recommendations"]:
                lines.append(f"- {r}")
        lines.append("")

    # 5. Critical gaps (items human SAP has but AI doesn't)
    lines.extend([
        "---",
        "",
        "## 5. Critical Gaps: Human SAP Content Missing from AI SAP",
        "",
        "These are areas where the human statistician provided content that the AI SAP lacks.",
        "These represent the highest-priority improvements.",
        "",
        "| Priority | Gap | Human SAP Has | AI SAP Status |",
        "|----------|-----|---------------|---------------|",
        "| P0 | Table shells (output templates) | ~70 tables | 0 tables |",
        "| P0 | FAS definition | Defined per convention | [SKIP] |",
        "| P0 | Baseline definition per parameter | Explicit for each parameter type | [SKIP] |",
        "| P1 | Protocol deviation classification | Defined criteria | Partial + [SKIP] |",
        "| P1 | Statistical software versions | SAS, WinNonlin versions | [SKIP] |",
        "| P1 | BLQ handling rules | Detailed (LLOQ/2, 0, etc.) | Present but less precise |",
        "| P1 | IP compliance formula | (Actual/Planned)×100 | Generic |",
        "| P1 | Missing data imputation rules | Detailed rules per type | Multiple [SKIP]s |",
        "| P2 | Screening test breakout | 4 separate sections | 1 combined |",
        "| P2 | Synopsis / Trial Schema / SoA | Full Section 3 | Not in template |",
        "| P2 | Dose proportionality formula | Power model equation | Referenced |",
        "| P2 | Preliminary PK deliverable format | File naming, templates | Generic |",
        "",
    ])

    # 6. Phase classification
    lines.extend([
        "---",
        "",
        "## 6. Improvement Phases",
        "",
        "### Phase 1 (Current): Protocol Extraction Accuracy",
        "Focus: Maximize content faithfully extracted from the protocol.",
        "- Reduce [SKIP] markers where protocol HAS the information",
        "- Improve specificity (extract exact numbers, formulas, versions)",
        "- Better BLQ handling, PK parameter, dose proportionality detail",
        "",
        "### Phase 2 (Next): Expert Judgment Layer",
        "Focus: Fill gaps where protocol is silent but convention is clear.",
        "- Generate standard FAS/PPS definitions using biostatistics conventions",
        "- Generate baseline definitions per parameter type",
        "- Generate missing data imputation rules (SAP-standard defaults)",
        "- Generate protocol deviation classification framework",
        "- Generate statistical software standard text",
        "",
        "### Phase 3 (Future): Table Shell Generation",
        "Focus: Generate output table templates (the #1 value-add of a SAP).",
        "- Table shells for demographics, disposition, AE summaries",
        "- Table shells for PK parameter summaries",
        "- Table shells for lab shift tables",
        "- Table shells for individual listings",
        "",
    ])

    return "\n".join(lines)


# ============================================================================
# Main
# ============================================================================

def main():
    parser = argparse.ArgumentParser(description="SAP Quality Verification")
    parser.add_argument("--ai-sap", required=True, help="Path to AI SAP content .txt file")
    parser.add_argument("--ai-docx", help="Path to AI SAP .docx file (for table counting)")
    parser.add_argument("--human-sap", required=True, help="Path to human SAP .docx file")
    parser.add_argument("--output", default="reports/verification_report.md")
    parser.add_argument("--skip-llm", action="store_true", help="Skip LLM evaluation (automated only)")
    parser.add_argument("--sections", nargs="*", help="Only evaluate these tags (default: all)")
    args = parser.parse_args()

    # Resolve AI docx path from content txt
    if not args.ai_docx:
        args.ai_docx = args.ai_sap.replace("_content.txt", ".docx")

    print("=" * 70)
    print("SAP QUALITY VERIFICATION")
    print("=" * 70)

    # Extract content
    print("\n[1/4] Extracting content...")
    ai_sections = extract_ai_sections(args.ai_sap)
    human_text = extract_human_sections(args.human_sap)
    print(f"  AI SAP: {len(ai_sections)} sections")
    print(f"  Human SAP: {len(human_text):,} chars")

    # Automated checks
    print("\n[2/4] Running automated checks...")
    auto_results = run_automated_checks(ai_sections, human_text, args.ai_docx, args.human_sap)
    skip_total = auto_results["skip_markers"].get("_total", 0)
    print(f"  [SKIP] markers: {skip_total}")
    print(f"  Tables: AI={auto_results['table_comparison']['ai_tables']}, "
          f"Human={auto_results['table_comparison']['human_tables']}")

    # Thin sections
    thin = auto_results["coverage"]["thin_sections"]
    if thin:
        print(f"  Thin sections (<30 words): {list(thin.keys())}")

    # LLM evaluation
    section_scores = []
    if not args.skip_llm:
        print("\n[3/4] Running LLM evaluation (this takes a few minutes)...")

        tags_to_eval = args.sections or [tag for tag, _, _, _ in SECTION_MAP]

        for i, (ai_tag, human_id, human_title, category) in enumerate(SECTION_MAP):
            if ai_tag not in tags_to_eval:
                section_scores.append({"tag": ai_tag, "summary": "skipped"})
                continue

            ai_text = ai_sections.get(ai_tag, "")
            if not ai_text:
                section_scores.append({
                    "tag": ai_tag,
                    "protocol_fidelity": 0,
                    "completeness": 0,
                    "specificity": 0,
                    "comparison_to_human": 0,
                    "gaps": ["Section entirely missing from AI SAP"],
                    "strengths": [],
                    "recommendations": ["Add this section"],
                    "summary": "Missing",
                })
                continue

            human_excerpt = find_human_excerpt(human_text, human_title, ai_tag)
            print(f"  [{i+1}/{len(SECTION_MAP)}] Evaluating {ai_tag}...", end=" ", flush=True)

            score = evaluate_section_pair(ai_tag, ai_text, human_excerpt, category, human_title)
            score["tag"] = ai_tag
            section_scores.append(score)
            f = score.get("protocol_fidelity", -1)
            c = score.get("comparison_to_human", -1)
            print(f"fidelity={f}, vs_human={c}")
    else:
        print("\n[3/4] LLM evaluation skipped (--skip-llm)")
        for ai_tag, _, _, _ in SECTION_MAP:
            section_scores.append({"tag": ai_tag, "summary": "not evaluated"})

    # Generate report
    print("\n[4/4] Generating report...")
    report = generate_report(auto_results, section_scores, args.ai_sap, args.human_sap)

    os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
    with open(args.output, "w") as f:
        f.write(report)
    print(f"\n  Report saved to: {args.output}")

    # Print quick summary
    scored = [s for s in section_scores if s.get("protocol_fidelity", -1) >= 0]
    if scored:
        avg = sum(s.get("comparison_to_human", 0) for s in scored) / len(scored)
        low = [s["tag"] for s in scored if s.get("comparison_to_human", 0) <= 2]
        print(f"\n  Average vs-human score: {avg:.1f}/5")
        if low:
            print(f"  Low-scoring sections: {low}")

    print("\nDone.")


if __name__ == "__main__":
    main()
