"""
Create Yuhan SAP Jinja2 template by reading the original SAP docx
and replacing body content with {{ tag }} placeholders.

Usage:
    source YH-SAP-venv/bin/activate
    python scripts/create_yuhan_template.py
"""

import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
INPUT_PATH = PROJECT_ROOT / "SBS02T01 Ver02_Statistical Analysis Plan_20251212.docx"
OUTPUT_PATH = PROJECT_ROOT / "sap-kcl" / "auto_sap" / "templates" / "yuhan_sap_template_v1.0.docx"


# ── Section heading → tag mapping ──────────────────────────────────────────
# Keys are (heading_level, heading_text_startswith) tuples.
# Values are the Jinja2 tag name to insert.
# None means "keep as-is" (don't template this section).

HEADING_TAG_MAP = {
    # Section 2
    (1, "INTRODUCTION"): "introduction",
    # Section 3
    (1, "OBJECTIVES"): "objectives",
    # Section 4 - Analysis Sets
    (1, "ANALYSIS SETS"): None,  # keep heading only, subsections handle body
    (2, "Screened Set"): "screened_set",
    (2, "Randomized Set"): "randomized_set",
    (2, "Safety Set"): "safety_set",
    (2, "Full Analysis Set"): "fas_definition",
    (2, "Per-Protocol Set"): "pps_definition",
    # Section 5
    (1, "SUBJECT DISPOSITION"): "subject_disposition",
    # Section 6
    (1, "DEMOGRAPHICS AND OTHER BASELINE"): "demographics",
    # Section 7
    (1, "EXTENT OF EXPOSURE"): None,
    (2, "Investigational Product"): "ip_exposure",
    (2, "Measurement of Treatment Compliance"): "treatment_compliance",
    # Section 8
    (1, "MEDICAL HISTORY"): None,
    (2, "Medical History"): "medical_history",
    (2, "Prior and Concomitant Medication"): "concomitant_medication",
    # Section 9
    (1, "EFFICACY ANALYSIS"): None,  # keep overview text as-is? No, tag it.
    (2, "Primary Efficacy Parameter"): "primary_efficacy",
    (2, "Secondary Efficacy Parameter"): "secondary_efficacy",
    (2, "Additional Efficacy"): "additional_efficacy",
    # Section 10
    (1, "SAFETY ANALYSIS"): None,
    (2, "Adverse Events"): "adverse_events",
    (2, "Clinical Laboratory Parameters"): "lab_parameters",
    (2, "Vital Signs"): "vital_signs",
    (2, "Electrocardiogram"): "ecg",
    (2, "Physical Examination"): "physical_exam",
    (2, "Other Safety Parameter"): "other_safety",
    # Section 11
    (1, "INTERIM ANALYSIS"): "interim_analysis",
    # DMC
    (1, "DMC ANALYSIS"): "dmc_analysis",
    # Section 12
    (1, "DETERMINATION OF SAMPLE SIZE"): "sample_size",
    # Section 13 - keep standard text as-is
    (1, "STATISTICAL SOFTWARE"): None,
    # Statistical Considerations
    (1, "STATISTICAL CONSIDERATIONS"): "statistical_considerations",
    # Section 14 - Data Handling
    (1, "DATA HANDLING CONVENTIONS"): None,
    (2, "REFERENCE START DATE"): None,  # keep as-is (standard)
    (2, "Visit Time Windows"): "visit_windows",
    (2, "Derived Efficacy Variables"): "derived_variables",
    (2, "Repeated or Unscheduled"): "repeated_assessments",
    (2, "Missing Date of the Last Dose"): "missing_last_dose_date",
    (2, "Missing Severity Assessment"): "missing_ae_severity",
    (2, "Missing Causal Relationship"): "missing_causality",
    (2, "Missing Date Information"): "missing_dates",
    (2, "Character Values"): "lab_char_values",
    # Section 15
    (1, "CHANGES TO ANALYSIS"): "protocol_changes",
    # Sections to skip entirely
    (1, "LIST OF ABBREVIATIONS"): None,
    (1, "REFERENCES"): None,
    (1, "APPENDIX"): None,
}


def get_heading_level(para):
    """Return heading level (1, 2, 3) or 0 if not a heading."""
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 0
    return 0


def match_heading(para):
    """Try to match a paragraph to a heading tag map entry. Returns tag or None."""
    level = get_heading_level(para)
    if level == 0:
        return None, False

    text = para.text.strip()
    for (hlevel, prefix), tag in HEADING_TAG_MAP.items():
        if level == hlevel and text.upper().startswith(prefix.upper()):
            return tag, True
    return None, False


def is_body_paragraph(para):
    """Check if a paragraph is body content (not heading, not TOC, not empty)."""
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        return False
    if style_name.startswith("toc ") or style_name == "table of figures":
        return False
    return True


def clear_paragraph_content(para):
    """Remove all runs from a paragraph, keeping the paragraph element."""
    for run in para.runs:
        run._element.getparent().remove(run._element)
    # Also remove any remaining text nodes
    for child in list(para._element):
        if child.tag == qn("w:r"):
            para._element.remove(child)


def set_paragraph_text(para, text):
    """Set paragraph text, preserving the paragraph's style."""
    clear_paragraph_content(para)
    run = para.add_run(text)
    return run


def process_document(doc):
    """Walk through paragraphs, replacing body content with Jinja2 tags."""
    paragraphs = doc.paragraphs
    n = len(paragraphs)

    # Build a list of (start_idx, end_idx, tag) for sections to replace
    sections_to_tag = []

    i = 0
    while i < n:
        tag, is_heading = match_heading(paragraphs[i])

        if is_heading and tag is not None:
            # Found a heading that needs tagging
            heading_idx = i
            # Find the body paragraphs after this heading
            body_start = i + 1
            body_end = body_start

            # Scan forward until next heading of same or higher level, or end
            heading_level = get_heading_level(paragraphs[i])
            j = body_start
            while j < n:
                next_level = get_heading_level(paragraphs[j])
                if next_level > 0 and next_level <= heading_level:
                    break
                # Also stop at next heading of any level in the map
                _, next_is_heading = match_heading(paragraphs[j])
                if next_is_heading:
                    break
                j += 1
            body_end = j

            sections_to_tag.append((body_start, body_end, tag))
            i = body_end
        else:
            i += 1

    # Process sections in reverse order to preserve indices
    for body_start, body_end, tag in reversed(sections_to_tag):
        if body_start >= body_end:
            # No body paragraphs — need to insert a new paragraph after the heading
            # The heading is at body_start - 1
            heading_para = paragraphs[body_start - 1]
            new_p = copy.deepcopy(heading_para._element)
            # Clear formatting to use body style, then set tag text
            from docx.oxml import OxmlElement
            new_para_elem = OxmlElement("w:p")
            heading_para._element.addnext(new_para_elem)
            # Now we need to add the tag text via a run
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = "{{ " + tag + " }}"
            new_r.append(new_t)
            new_para_elem.append(new_r)
            continue

        # Keep the first body paragraph, set its text to the tag
        first_body = paragraphs[body_start]

        # Remove all body paragraphs after the first one
        for idx in range(body_end - 1, body_start, -1):
            p = paragraphs[idx]
            p_element = p._element
            p_element.getparent().remove(p_element)

        # Set the first body paragraph to the Jinja2 tag
        set_paragraph_text(first_body, "{{ " + tag + " }}")

    return doc


def add_efficacy_overview_tag(doc):
    """Special handling: add efficacy overview tag after EFFICACY ANALYSIS heading."""
    for i, para in enumerate(doc.paragraphs):
        if get_heading_level(para) == 1 and para.text.strip().startswith("EFFICACY ANALYSIS"):
            # Find the body paragraph(s) between this heading and the next H2
            j = i + 1
            while j < len(doc.paragraphs):
                if doc.paragraphs[j].text.strip():
                    level = get_heading_level(doc.paragraphs[j])
                    if level > 0:
                        break
                    # This is the overview paragraph - don't replace, keep it
                    # Actually, check if it has content
                    text = doc.paragraphs[j].text.strip()
                    if text and not text.startswith("{{"):
                        # There's overview text before the subsections
                        # We already handled it in the main pass if the heading had a tag
                        pass
                    break
                j += 1
            break


def add_statistical_software_version_tag(doc):
    """In the STATISTICAL SOFTWARE section, add version tag where version number appears."""
    for i, para in enumerate(doc.paragraphs):
        if get_heading_level(para) == 1 and para.text.strip().startswith("STATISTICAL SOFTWARE"):
            # Find the body paragraph
            j = i + 1
            while j < len(doc.paragraphs):
                text = doc.paragraphs[j].text.strip()
                if text:
                    if "9.4" in text or "version" in text.lower():
                        # Replace "version 9.4 (or later)" with tag
                        new_text = text.replace(
                            "version 9.4 (or later)",
                            "version {{ statistical_software_version }}"
                        )
                        if new_text == text:
                            # Try alternative patterns
                            new_text = text.replace("9.4", "{{ statistical_software_version }}")
                        set_paragraph_text(doc.paragraphs[j], new_text)
                    break
                j += 1
            break


def main():
    print(f"Reading: {INPUT_PATH}")
    doc = Document(str(INPUT_PATH))

    print("Processing sections and inserting Jinja2 tags...")
    process_document(doc)

    # Special handling for statistical software version
    add_statistical_software_version_tag(doc)

    # Ensure output directory exists
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    print(f"Saving: {OUTPUT_PATH}")
    doc.save(str(OUTPUT_PATH))

    # Verify by listing tags
    print("\nVerifying tags in output document:")
    verify_doc = Document(str(OUTPUT_PATH))
    tag_count = 0
    for p in verify_doc.paragraphs:
        text = p.text.strip()
        if "{{" in text and "}}" in text:
            print(f"  ✓ {text}")
            tag_count += 1
    print(f"\nTotal tags found: {tag_count}")


if __name__ == "__main__":
    main()
